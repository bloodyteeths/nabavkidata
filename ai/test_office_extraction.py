#!/usr/bin/env python3
"""
Test script for Office document extraction
Creates sample documents with realistic tender data and tests extraction
"""

import os
import sys
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add parent directory to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from ai.office_extraction import extract_from_office_document


def create_test_excel_boq(filepath: str):
    """Create a sample Excel Bill of Quantities in Macedonian"""
    print(f"Creating test Excel file: {filepath}")

    # Sample BOQ data in Macedonian
    data = {
        'Р.Бр.': [1, 2, 3, 4, 5, 6, 7, 8],
        'Назив на производот': [
            'Компјутер Lenovo ThinkCentre M720',
            'Монитор Dell 24 инчи',
            'Тастатура и глушец',
            'UPS 1000VA',
            'Принтер HP LaserJet Pro',
            'Скенер Canon',
            'Мрежен свич 24-port',
            'Кабли и конектори'
        ],
        'Количина': [50, 100, 50, 25, 10, 5, 3, 150],
        'Единица мера': ['парче', 'парче', 'сет', 'парче', 'парче', 'парче', 'парче', 'парче'],
        'Единечна цена (МКД)': [35000, 12000, 1500, 8000, 18000, 15000, 25000, 500],
        'Вкупна цена (МКД)': [1750000, 1200000, 75000, 200000, 180000, 75000, 75000, 75000],
        'Технички карактеристики': [
            'Intel i5, 8GB RAM, 256GB SSD',
            'Full HD, IPS панел',
            'USB, безжични',
            'Line-interactive, 10 min backup',
            'A4, duplex, network',
            'Автоматски податоци, ADF',
            'Gigabit Ethernet, managed',
            'CAT6, RJ45'
        ]
    }

    df = pd.DataFrame(data)

    # Create Excel file with multiple sheets
    with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Предмер и предрачка', index=False)

        # Add a second sheet with bidder info
        bidder_data = {
            'Понудувач': ['ТехноМарт ДООЕЛ', 'КомпТех АД', 'ИТ Солушнс'],
            'Даночен број': ['4057023501234', '4057023502345', '4057023503456'],
            'Вкупна понуда (МКД)': [3630000, 3580000, 3720000],
            'Рок на испорака (дена)': [30, 45, 25]
        }
        pd.DataFrame(bidder_data).to_excel(writer, sheet_name='Понудувачи', index=False)

    print(f"Created Excel file with {len(df)} items")


def create_test_word_table(filepath: str):
    """Create a sample Word document with table in Macedonian"""
    print(f"Creating test Word file: {filepath}")

    doc = Document()

    # Add title
    title = doc.add_heading('Образец за техничка спецификација', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add some text
    doc.add_paragraph(
        'Предмет на набавка: Набавка на канцелариски материјали за 2025 година'
    )
    doc.add_paragraph(
        'Тендер број: 21234/2025'
    )
    doc.add_paragraph('')

    # Add table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    headers = ['Р.Бр.', 'Опис', 'Количина', 'Единица', 'Единечна цена', 'Вкупна цена']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True

    # Data rows
    items = [
        ('1', 'Хартија А4 80g/m² - бела', '500', 'пакети', '250.00', '125,000.00'),
        ('2', 'Фасцикли картонски со кукичка', '200', 'парче', '30.00', '6,000.00'),
        ('3', 'Пенкала сина', '300', 'парче', '15.00', '4,500.00'),
        ('4', 'Маркери за табла - различни бои', '100', 'парче', '80.00', '8,000.00'),
        ('5', 'Тонер за принтер HP LaserJet', '50', 'парче', '3,500.00', '175,000.00'),
        ('6', 'Стиплер метален', '25', 'парче', '150.00', '3,750.00'),
        ('7', 'Спојници за документи', '100', 'кутија', '120.00', '12,000.00'),
    ]

    for item_data in items:
        row = table.add_row()
        for i, value in enumerate(item_data):
            row.cells[i].text = value

    # Add footer paragraph
    doc.add_paragraph('')
    doc.add_paragraph('Вкупна вредност: 334,250.00 МКД')
    doc.add_paragraph('Напомена: Цените се без ДДВ.')

    doc.save(filepath)
    print(f"Created Word file with {len(items)} items in table")


def test_extraction(filepath: str):
    """Test extraction on a file"""
    print(f"\n{'='*70}")
    print(f"TESTING EXTRACTION: {os.path.basename(filepath)}")
    print(f"{'='*70}\n")

    result = extract_from_office_document(filepath)

    # Print summary
    print(f"Results:")
    print(f"  File Type: {result['metadata']['file_type']}")
    print(f"  Tables Found: {result['metadata']['table_count']}")
    print(f"  Items Extracted: {result['metadata']['item_count']}")
    print(f"  Text Length: {result['metadata']['text_length']}")

    # Print table details
    if result['tables']:
        print(f"\n{'='*70}")
        print(f"TABLE DETAILS")
        print(f"{'='*70}")
        for i, table in enumerate(result['tables'], 1):
            print(f"\nTable {i}:")
            print(f"  Type: {table.table_type}")
            print(f"  Source: {table.source_type}")
            print(f"  Sheet: {table.sheet_name if table.sheet_name else 'N/A'}")
            print(f"  Dimensions: {table.metadata['rows']} rows × {table.metadata['cols']} columns")
            print(f"  Confidence: {table.confidence}")
            print(f"  Columns: {list(table.data.columns)[:5]}...")  # First 5 columns

    # Print item details
    if result['items']:
        print(f"\n{'='*70}")
        print(f"EXTRACTED ITEMS (showing first 5)")
        print(f"{'='*70}")
        for i, item in enumerate(result['items'][:5], 1):
            print(f"\nItem {i}:")
            print(f"  Number: {item['item_number']}")
            print(f"  Name: {item['item_name']}")
            print(f"  Quantity: {item['quantity']} {item['unit']}")
            print(f"  Unit Price: {item['unit_price']:,.2f} МКД" if item['unit_price'] else "  Unit Price: N/A")
            print(f"  Total Price: {item['total_price']:,.2f} МКД" if item['total_price'] else "  Total Price: N/A")
            print(f"  Specs: {item['specifications'][:50]}..." if item['specifications'] else "  Specs: N/A")
            print(f"  Confidence: {item['confidence']}")

        if len(result['items']) > 5:
            print(f"\n... and {len(result['items']) - 5} more items")

        # Calculate total value
        total_value = sum(item['total_price'] for item in result['items'] if item['total_price'])
        print(f"\nTotal Extracted Value: {total_value:,.2f} МКД")

    # Print sample text (for Word docs)
    if result['text']:
        print(f"\n{'='*70}")
        print(f"DOCUMENT TEXT (first 300 chars)")
        print(f"{'='*70}")
        print(result['text'][:300])
        if len(result['text']) > 300:
            print("...")

    print(f"\n{'='*70}\n")

    return result


def main():
    # Create test directory
    test_dir = '/Users/tamsar/Downloads/nabavkidata/ai/test_documents'
    os.makedirs(test_dir, exist_ok=True)

    # Create test files
    excel_file = os.path.join(test_dir, 'test_boq_macedonian.xlsx')
    word_file = os.path.join(test_dir, 'test_specs_macedonian.docx')

    print("="*70)
    print("OFFICE DOCUMENT EXTRACTION TEST")
    print("="*70)
    print()

    # Create test Excel BOQ
    create_test_excel_boq(excel_file)

    # Create test Word document
    create_test_word_table(word_file)

    print()

    # Test extraction on Excel
    excel_result = test_extraction(excel_file)

    # Test extraction on Word
    word_result = test_extraction(word_file)

    # Final summary
    print("="*70)
    print("FINAL TEST SUMMARY")
    print("="*70)
    print(f"Excel File:")
    print(f"  Tables: {excel_result['metadata']['table_count']}")
    print(f"  Items: {excel_result['metadata']['item_count']}")
    print()
    print(f"Word File:")
    print(f"  Tables: {word_result['metadata']['table_count']}")
    print(f"  Items: {word_result['metadata']['item_count']}")
    print()
    print(f"Total Items Extracted: {excel_result['metadata']['item_count'] + word_result['metadata']['item_count']}")
    print("="*70)


if __name__ == '__main__':
    main()
