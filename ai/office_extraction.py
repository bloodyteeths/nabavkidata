"""
Office Document Extraction Pipeline for Nabavkidata
Handles: DOCX, DOC, XLSX, XLS

Key capabilities:
- Extract all tables from Word documents
- Extract all sheets from Excel documents
- Parse Macedonian column headers
- Identify BOQ/price list/spec tables
- Extract structured item data

Author: AI Assistant
Date: 2025-11-29
"""

import os
import re
import logging
from typing import List, Dict, Optional, Tuple, Any
from dataclasses import dataclass, asdict
import pandas as pd
import json

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class ExtractedOfficeTable:
    """Represents a table extracted from an Office document"""
    data: pd.DataFrame
    source_type: str  # 'word' or 'excel'
    sheet_name: Optional[str]  # For Excel
    table_index: int
    confidence: float
    table_type: str  # items, specs, bidders, pricing, evaluation
    metadata: Dict[str, Any]

    def to_dict(self) -> Dict:
        """Convert to dictionary (excluding DataFrame)"""
        d = asdict(self)
        d['data'] = self.data.to_dict('records')
        return d


class WordExtractor:
    """Extract tables and text from Word documents"""

    def __init__(self):
        try:
            from docx import Document
            self.Document = Document
            self.available = True
            logger.info("WordExtractor initialized successfully")
        except ImportError as e:
            self.available = False
            logger.warning(f"python-docx not available: {e}")

    def extract_tables(self, doc_path: str) -> List[ExtractedOfficeTable]:
        """Extract all tables from a Word document"""
        if not self.available:
            logger.error("WordExtractor not available - python-docx not installed")
            return []

        if not os.path.exists(doc_path):
            logger.error(f"Document not found: {doc_path}")
            return []

        tables = []

        try:
            doc = self.Document(doc_path)
            logger.info(f"Processing {len(doc.tables)} tables from {os.path.basename(doc_path)}")

            for idx, table in enumerate(doc.tables):
                # Convert table to DataFrame
                data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    data.append(row_data)

                if len(data) > 1:  # Has header + data
                    # Use first row as column names
                    df = pd.DataFrame(data[1:], columns=data[0])

                    # Skip empty tables
                    if df.empty or df.dropna(how='all').empty:
                        continue

                    table_type = self._classify_table(df)

                    tables.append(ExtractedOfficeTable(
                        data=df,
                        source_type='word',
                        sheet_name=None,
                        table_index=idx,
                        confidence=0.85,
                        table_type=table_type,
                        metadata={
                            'rows': len(df),
                            'cols': len(df.columns),
                            'source_file': os.path.basename(doc_path)
                        }
                    ))

                    logger.info(f"Extracted table {idx}: {table_type} ({len(df)} rows, {len(df.columns)} cols)")

            logger.info(f"Total tables extracted from Word: {len(tables)}")

        except Exception as e:
            logger.error(f"Error extracting tables from Word document {doc_path}: {e}", exc_info=True)

        return tables

    def extract_text(self, doc_path: str) -> str:
        """Extract full text from Word document"""
        if not self.available:
            return ""

        if not os.path.exists(doc_path):
            logger.error(f"Document not found: {doc_path}")
            return ""

        try:
            doc = self.Document(doc_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            logger.info(f"Extracted {len(text)} characters of text from {os.path.basename(doc_path)}")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from {doc_path}: {e}")
            return ""

    def _classify_table(self, df: pd.DataFrame) -> str:
        """Classify table type based on columns"""
        cols_lower = ' '.join(str(c) for c in df.columns).lower()

        # Macedonian patterns for BOQ/Items
        if any(p in cols_lower for p in ['количина', 'количество', 'кол.']):
            return 'items'
        if any(p in cols_lower for p in ['цена', 'единечна цена', 'вкупна цена']):
            return 'items'
        if any(p in cols_lower for p in ['назив', 'опис', 'артикл', 'производ']):
            return 'items'
        if any(p in cols_lower for p in ['единица', 'мерка', 'ед. мера', 'јединица']):
            return 'items'

        # Macedonian patterns for Bidders
        if any(p in cols_lower for p in ['понудувач', 'понудувачи', 'фирма', 'компанија']):
            return 'bidders'
        if any(p in cols_lower for p in ['еднорни', 'даночен број', 'емб']):
            return 'bidders'

        # Macedonian patterns for Specifications
        if any(p in cols_lower for p in ['спецификација', 'технички', 'барања', 'карактеристики']):
            return 'specifications'

        # English patterns for Items
        if any(p in cols_lower for p in ['quantity', 'qty', 'amount']):
            return 'items'
        if any(p in cols_lower for p in ['price', 'unit price', 'total']):
            return 'items'
        if any(p in cols_lower for p in ['description', 'item', 'product', 'article']):
            return 'items'
        if any(p in cols_lower for p in ['unit', 'measure', 'uom']):
            return 'items'

        # English patterns for Bidders
        if any(p in cols_lower for p in ['bidder', 'vendor', 'supplier', 'company']):
            return 'bidders'

        # Check if it looks like an evaluation matrix
        if any(p in cols_lower for p in ['оценка', 'бодови', 'рангирање', 'score', 'rating']):
            return 'evaluation'

        return 'unknown'


class ExcelExtractor:
    """Extract tables from Excel documents"""

    def __init__(self):
        self.available = True
        try:
            import openpyxl
            import xlrd
            logger.info("ExcelExtractor initialized successfully")
        except ImportError as e:
            logger.warning(f"Excel library missing: {e}")
            self.available = False

    def extract_tables(self, excel_path: str) -> List[ExtractedOfficeTable]:
        """Extract all sheets as tables from Excel"""
        if not self.available:
            logger.error("ExcelExtractor not available")
            return []

        if not os.path.exists(excel_path):
            logger.error(f"Excel file not found: {excel_path}")
            return []

        tables = []

        try:
            # Try reading with pandas (handles both xlsx and xls)
            # For .xls files, we need xlrd
            ext = os.path.splitext(excel_path)[1].lower()
            engine = 'xlrd' if ext == '.xls' else 'openpyxl'

            xl = pd.ExcelFile(excel_path, engine=engine)
            logger.info(f"Processing {len(xl.sheet_names)} sheets from {os.path.basename(excel_path)}")

            for idx, sheet_name in enumerate(xl.sheet_names):
                try:
                    df = pd.read_excel(xl, sheet_name=sheet_name)

                    # Skip empty sheets
                    if df.empty or len(df) < 2:
                        logger.debug(f"Skipping empty sheet: {sheet_name}")
                        continue

                    # Clean the dataframe
                    df = self._clean_dataframe(df)

                    if df.empty:
                        continue

                    table_type = self._classify_table(df, sheet_name)

                    tables.append(ExtractedOfficeTable(
                        data=df,
                        source_type='excel',
                        sheet_name=sheet_name,
                        table_index=idx,
                        confidence=0.90,  # Excel is usually well-structured
                        table_type=table_type,
                        metadata={
                            'rows': len(df),
                            'cols': len(df.columns),
                            'sheet': sheet_name,
                            'source_file': os.path.basename(excel_path)
                        }
                    ))

                    logger.info(f"Extracted sheet '{sheet_name}': {table_type} ({len(df)} rows, {len(df.columns)} cols)")

                except Exception as e:
                    logger.warning(f"Error processing sheet {sheet_name}: {e}")
                    continue

            logger.info(f"Total tables extracted from Excel: {len(tables)}")

        except Exception as e:
            logger.error(f"Error extracting Excel {excel_path}: {e}", exc_info=True)

        return tables

    def _clean_dataframe(self, df: pd.DataFrame) -> pd.DataFrame:
        """Clean up Excel dataframe"""
        # Drop completely empty rows/columns
        df = df.dropna(how='all').dropna(axis=1, how='all')

        if df.empty:
            return df

        # If first row looks like headers but isn't set, fix it
        if df.columns[0] == 0 or 'Unnamed' in str(df.columns[0]):
            # Check if first row contains non-numeric data (likely headers)
            first_row = df.iloc[0]
            if first_row.apply(lambda x: isinstance(x, str)).sum() > len(first_row) / 2:
                df.columns = df.iloc[0]
                df = df.iloc[1:]

        # Reset index
        df = df.reset_index(drop=True)

        # Remove rows that are all NaN after cleaning
        df = df.dropna(how='all')

        return df

    def _classify_table(self, df: pd.DataFrame, sheet_name: str) -> str:
        """Classify Excel table type"""
        cols_str = ' '.join(str(c) for c in df.columns).lower()
        sheet_lower = sheet_name.lower()

        # Check sheet name first
        if any(p in sheet_lower for p in ['предмер', 'stavki', 'artikli', 'items', 'boq']):
            return 'items'
        if any(p in sheet_lower for p in ['понуда', 'понудувач', 'bidder', 'vendor']):
            return 'bidders'

        # BOQ/Items patterns (Macedonian)
        if any(p in cols_str for p in ['р.бр', 'р. бр', 'реден', 'ред. број']):
            return 'items'
        if any(p in cols_str for p in ['назив', 'опис', 'артикл', 'производ']):
            return 'items'
        if any(p in cols_str for p in ['количина', 'кол.', 'количество']):
            return 'items'
        if any(p in cols_str for p in ['единечна цена', 'ед. цена', 'цена']):
            return 'items'
        if any(p in cols_str for p in ['вкупна цена', 'вкупно', 'износ']):
            return 'items'
        if any(p in cols_str for p in ['единица', 'мерка', 'ед. мера']):
            return 'items'

        # Bidder patterns
        if any(p in cols_str for p in ['понудувач', 'фирма', 'компанија']):
            return 'bidders'
        if any(p in cols_str for p in ['даночен', 'емб', 'еднорни']):
            return 'bidders'

        # English patterns
        if any(p in cols_str for p in ['item', 'no.', '#', 'line']):
            return 'items'
        if any(p in cols_str for p in ['description', 'name', 'product']):
            return 'items'
        if any(p in cols_str for p in ['qty', 'quantity']):
            return 'items'
        if any(p in cols_str for p in ['unit price', 'price', 'rate']):
            return 'items'
        if any(p in cols_str for p in ['total', 'amount', 'value']):
            return 'items'
        if any(p in cols_str for p in ['unit', 'uom', 'measure']):
            return 'items'

        return 'unknown'


class OfficeItemExtractor:
    """Extract structured items from Office tables"""

    # Column mapping patterns (regex -> standard name)
    COLUMN_PATTERNS = {
        # Macedonian patterns
        r'р\.?\s*бр|реден|ред\.?\s*број|ред\.?\s*бр': 'item_number',
        r'назив|име(?!\s*на)|опис|производ|артикл|предмет': 'item_name',
        r'^количина$|^кол\.?$|количество': 'quantity',
        r'единица|мерка|мера|ед\.?\s*мера|јед\.?\s*мера|јм': 'unit',
        r'единечна.*цена|цена.*единица|ед\.?\s*цена|единична\s*цена': 'unit_price',
        r'вкупн|тотал|укупно|вкупна.*цена|износ|total': 'total_price',
        r'спец|тех.*барањ|карактер|технички': 'specifications',

        # English patterns
        r'^item$|^no\.?$|^#$|^line$|item\s*no': 'item_number',
        r'desc|name|product|article|subject': 'item_name',
        r'^qty$|quantity|amount(?!.*price)': 'quantity',
        r'^unit$|measure|uom': 'unit',
        r'unit.*price|price.*unit|rate|price\s*per': 'unit_price',
        r'^total$|sum|value|amount|total.*price': 'total_price',
        r'spec|requirement|characteristic': 'specifications',
    }

    def extract_items(self, tables: List[ExtractedOfficeTable]) -> List[Dict]:
        """Extract items from tables"""
        all_items = []

        for table in tables:
            # Only process tables that might contain items
            if table.table_type in ['items', 'unknown']:
                try:
                    mapped_df = self._map_columns(table.data)
                    items = self._parse_items(mapped_df, table)
                    all_items.extend(items)

                    if items:
                        logger.info(f"Extracted {len(items)} items from {table.source_type} table {table.table_index}")

                except Exception as e:
                    logger.error(f"Error extracting items from table: {e}")

        return all_items

    def _map_columns(self, df: pd.DataFrame) -> pd.DataFrame:
        """Map columns to standard names"""
        column_map = {}

        for col in df.columns:
            col_str = str(col).lower().strip()

            for pattern, standard_name in self.COLUMN_PATTERNS.items():
                if re.search(pattern, col_str, re.IGNORECASE):
                    # Avoid duplicate mappings
                    if standard_name not in column_map.values():
                        column_map[col] = standard_name
                        logger.debug(f"Mapped column '{col}' -> '{standard_name}'")
                    break

        if column_map:
            df = df.rename(columns=column_map)

        return df

    def _parse_items(self, df: pd.DataFrame, table: ExtractedOfficeTable) -> List[Dict]:
        """Parse items from mapped dataframe"""
        items = []

        for idx, row in df.iterrows():
            item = {
                'item_name': self._get_value(row, 'item_name'),
                'item_number': self._get_value(row, 'item_number'),
                'quantity': self._parse_number(self._get_value(row, 'quantity')),
                'unit': self._get_value(row, 'unit'),
                'unit_price': self._parse_number(self._get_value(row, 'unit_price')),
                'total_price': self._parse_number(self._get_value(row, 'total_price')),
                'specifications': self._get_value(row, 'specifications'),
                'extraction_source': table.source_type,
                'source_sheet': table.sheet_name,
                'table_index': table.table_index,
                'confidence': table.confidence,
                'table_type': table.table_type,
            }

            # Only add if has meaningful data
            # Must have an item name and at least one other field
            if item['item_name'] and len(str(item['item_name'])) > 2:
                # Check if row has at least one numeric value
                if any([item['quantity'], item['unit_price'], item['total_price']]):
                    items.append(item)

        return items

    def _get_value(self, row, col_name: str) -> Optional[str]:
        """Safely get value from row"""
        if col_name in row.index:
            val = row[col_name]
            if pd.notna(val):
                return str(val).strip()
        return None

    def _parse_number(self, value: Optional[str]) -> Optional[float]:
        """Parse number from string (handles Macedonian format)"""
        if not value:
            return None

        try:
            # Remove currency symbols and text
            clean = re.sub(r'[МКДденEURденарденари€$\s]', '', str(value))

            # Remove any non-numeric characters except . and ,
            clean = re.sub(r'[^\d.,\-]', '', clean)

            if not clean:
                return None

            # Handle Macedonian format (1.234,56) vs English format (1,234.56)
            if ',' in clean and '.' in clean:
                # Determine which is the decimal separator
                last_comma = clean.rindex(',')
                last_dot = clean.rindex('.')

                if last_comma > last_dot:
                    # Comma is decimal: 1.234,56
                    clean = clean.replace('.', '').replace(',', '.')
                else:
                    # Dot is decimal: 1,234.56
                    clean = clean.replace(',', '')
            elif ',' in clean:
                # Could be either thousands or decimal
                # If it appears in last 3 positions, it's probably decimal
                if len(clean) - clean.rindex(',') <= 3:
                    clean = clean.replace(',', '.')
                else:
                    clean = clean.replace(',', '')

            return float(clean)
        except Exception as e:
            logger.debug(f"Could not parse number '{value}': {e}")
            return None


# Main extraction function
def extract_from_office_document(file_path: str) -> Dict:
    """
    Extract tables and items from Word or Excel document

    Args:
        file_path: Absolute path to the document

    Returns:
        {
            'tables': List[ExtractedOfficeTable],
            'items': List[Dict],
            'text': str (for Word only),
            'metadata': Dict
        }
    """
    ext = os.path.splitext(file_path)[1].lower()

    result = {
        'tables': [],
        'items': [],
        'text': '',
        'metadata': {
            'file_type': ext,
            'file_path': file_path,
            'file_name': os.path.basename(file_path)
        }
    }

    logger.info(f"Processing {ext} file: {os.path.basename(file_path)}")

    if ext in ['.docx', '.doc']:
        extractor = WordExtractor()
        if extractor.available:
            result['tables'] = extractor.extract_tables(file_path)
            result['text'] = extractor.extract_text(file_path)
        else:
            logger.error("WordExtractor not available")

    elif ext in ['.xlsx', '.xls']:
        extractor = ExcelExtractor()
        if extractor.available:
            result['tables'] = extractor.extract_tables(file_path)
        else:
            logger.error("ExcelExtractor not available")

    else:
        logger.warning(f"Unsupported file type: {ext}")

    # Extract items from tables
    if result['tables']:
        item_extractor = OfficeItemExtractor()
        result['items'] = item_extractor.extract_items(result['tables'])

    result['metadata']['table_count'] = len(result['tables'])
    result['metadata']['item_count'] = len(result['items'])
    result['metadata']['text_length'] = len(result['text'])

    logger.info(f"Extraction complete: {result['metadata']['table_count']} tables, {result['metadata']['item_count']} items")

    return result


if __name__ == "__main__":
    # Test the extractor
    import sys

    if len(sys.argv) > 1:
        test_file = sys.argv[1]
        print(f"\n{'='*60}")
        print(f"Testing Office Extraction on: {test_file}")
        print(f"{'='*60}\n")

        result = extract_from_office_document(test_file)

        print(f"\nResults:")
        print(f"  Tables found: {result['metadata']['table_count']}")
        print(f"  Items extracted: {result['metadata']['item_count']}")
        print(f"  Text length: {result['metadata']['text_length']}")

        print(f"\n{'='*60}")
        print(f"Table Details:")
        print(f"{'='*60}")
        for i, table in enumerate(result['tables']):
            print(f"\nTable {i+1}:")
            print(f"  Type: {table.table_type}")
            print(f"  Source: {table.source_type}")
            print(f"  Sheet: {table.sheet_name}")
            print(f"  Rows: {table.metadata['rows']}")
            print(f"  Columns: {table.metadata['cols']}")
            print(f"  Confidence: {table.confidence}")
            print(f"  Column names: {list(table.data.columns)}")

        if result['items']:
            print(f"\n{'='*60}")
            print(f"Sample Items (first 5):")
            print(f"{'='*60}")
            for i, item in enumerate(result['items'][:5]):
                print(f"\nItem {i+1}:")
                print(f"  Name: {item['item_name']}")
                print(f"  Number: {item['item_number']}")
                print(f"  Quantity: {item['quantity']}")
                print(f"  Unit: {item['unit']}")
                print(f"  Unit Price: {item['unit_price']}")
                print(f"  Total Price: {item['total_price']}")

        print(f"\n{'='*60}\n")
    else:
        print("Usage: python office_extraction.py <path_to_document>")
