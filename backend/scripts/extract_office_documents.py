#!/usr/bin/env python3
"""
Extract text from Word/Excel documents that failed initial extraction.
Downloads from e-nabavki.gov.mk and uses python-docx/openpyxl for extraction.
"""

import os
import sys
import tempfile
import logging
import time
import requests
from typing import Optional

# Add parent to path
sys.path.insert(0, '/Users/tamsar/Downloads/nabavkidata')
sys.path.insert(0, '/Users/tamsar/Downloads/nabavkidata/ai')

import psycopg2
from psycopg2.extras import RealDictCursor

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Database config
DB_CONFIG = {
    'host': 'nabavkidata-db.cb6gi2cae02j.eu-central-1.rds.amazonaws.com',
    'port': 5432,
    'user': 'nabavki_user',
    'password': os.getenv('DB_PASSWORD', ''),
    'database': 'nabavkidata',
}


def extract_word_text(file_path: str) -> Optional[str]:
    """Extract text from Word document"""
    try:
        from docx import Document
        doc = Document(file_path)

        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return '\n'.join(text_parts)
    except Exception as e:
        logger.error(f"Word extraction failed: {e}")
        return None


def extract_excel_text(file_path: str) -> Optional[str]:
    """Extract text from Excel document"""
    try:
        import pandas as pd

        # Determine engine
        ext = os.path.splitext(file_path)[1].lower()
        engine = 'xlrd' if ext == '.xls' else 'openpyxl'

        xl = pd.ExcelFile(file_path, engine=engine)

        text_parts = []
        for sheet_name in xl.sheet_names:
            try:
                df = pd.read_excel(xl, sheet_name=sheet_name)
                if not df.empty:
                    text_parts.append(f"=== Sheet: {sheet_name} ===")
                    # Convert to string representation
                    text_parts.append(df.to_string())
            except Exception as e:
                logger.warning(f"Error reading sheet {sheet_name}: {e}")

        return '\n\n'.join(text_parts)
    except Exception as e:
        logger.error(f"Excel extraction failed: {e}")
        return None


def download_document(url: str, file_name: str) -> Optional[str]:
    """Download document to temp file, return path"""
    try:
        response = requests.get(url, timeout=30)
        response.raise_for_status()

        # Get extension from filename
        ext = os.path.splitext(file_name)[1].lower()
        if not ext:
            ext = '.doc'  # Default

        # Save to temp file
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as f:
            f.write(response.content)
            return f.name
    except Exception as e:
        logger.error(f"Download failed: {e}")
        return None


def process_documents(limit: int = None, retry_failed: bool = True):
    """Process Word/Excel documents that need extraction"""

    conn = psycopg2.connect(
        host=DB_CONFIG['host'],
        port=DB_CONFIG['port'],
        user=DB_CONFIG['user'],
        password=DB_CONFIG['password'],
        dbname=DB_CONFIG['database']
    )
    cur = conn.cursor(cursor_factory=RealDictCursor)

    try:
        # Get documents to process
        statuses = ['pending', 'failed']
        if retry_failed:
            statuses.extend(['auth_required'])

        status_clause = ','.join(f"'{s}'" for s in statuses)

        query = f"""
            SELECT doc_id, file_url, file_name, tender_id, extraction_status
            FROM documents
            WHERE (file_name LIKE '%.doc%' OR file_name LIKE '%.xls%')
              AND extraction_status IN ({status_clause})
              AND file_url IS NOT NULL
              AND file_url != ''
            ORDER BY
                CASE WHEN file_name LIKE '%.xlsx' THEN 1
                     WHEN file_name LIKE '%.docx' THEN 2
                     WHEN file_name LIKE '%.xls' THEN 3
                     ELSE 4 END
        """
        if limit:
            query += f" LIMIT {limit}"

        cur.execute(query)
        docs = cur.fetchall()
        logger.info(f"Found {len(docs)} documents to process")

        success_count = 0
        failed_count = 0

        for i, doc in enumerate(docs):
            doc_id = doc['doc_id']
            file_url = doc['file_url']
            file_name = doc['file_name']
            tender_id = doc['tender_id']

            logger.info(f"[{i+1}/{len(docs)}] Processing: {file_name[:50]}...")

            # Download
            temp_path = download_document(file_url, file_name)
            if not temp_path:
                cur.execute(
                    "UPDATE documents SET extraction_status = 'download_failed' WHERE doc_id = %s",
                    (doc_id,)
                )
                conn.commit()
                failed_count += 1
                continue

            try:
                # Extract based on file type
                ext = os.path.splitext(file_name)[1].lower()

                if ext in ['.doc', '.docx']:
                    content = extract_word_text(temp_path)
                elif ext in ['.xls', '.xlsx']:
                    content = extract_excel_text(temp_path)
                else:
                    content = None

                if content and len(content) > 50:
                    # Update database
                    cur.execute("""
                        UPDATE documents
                        SET content_text = %s,
                            extraction_status = 'success',
                            updated_at = NOW()
                        WHERE doc_id = %s
                    """, (content[:500000], doc_id))  # Limit to 500k chars
                    conn.commit()

                    logger.info(f"  ✓ Extracted {len(content)} chars")
                    success_count += 1
                else:
                    cur.execute(
                        "UPDATE documents SET extraction_status = 'skip_empty' WHERE doc_id = %s",
                        (doc_id,)
                    )
                    conn.commit()
                    logger.warning(f"  ✗ No content extracted")
                    failed_count += 1

            except Exception as e:
                logger.error(f"  ✗ Extraction error: {e}")
                cur.execute(
                    "UPDATE documents SET extraction_status = 'failed' WHERE doc_id = %s",
                    (doc_id,)
                )
                conn.commit()
                failed_count += 1
            finally:
                # Clean up temp file
                if temp_path and os.path.exists(temp_path):
                    os.unlink(temp_path)

            # Rate limit
            time.sleep(0.5)

        logger.info(f"\n{'='*50}")
        logger.info(f"EXTRACTION COMPLETE")
        logger.info(f"{'='*50}")
        logger.info(f"Success: {success_count}")
        logger.info(f"Failed: {failed_count}")
        logger.info(f"Total: {len(docs)}")

    finally:
        cur.close()
        conn.close()


def main():
    import argparse
    parser = argparse.ArgumentParser(description='Extract Word/Excel documents')
    parser.add_argument('--limit', type=int, default=None, help='Max documents to process')
    parser.add_argument('--retry-failed', action='store_true', default=True, help='Retry failed documents')
    args = parser.parse_args()

    process_documents(limit=args.limit, retry_failed=args.retry_failed)


if __name__ == '__main__':
    main()
